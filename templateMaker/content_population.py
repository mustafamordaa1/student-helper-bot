import itertools
import random
import docx
from config import WORD_MAIN_PATH
from utils import database
from docxtpl import DocxTemplate

PLACEHOLDERS = {
    "الجواب الصحيح": "Correct Answer",
    "اسم الصورة": "Image Placeholder",
    "نص السؤال مدقق": "Question Text",
    "الخيار أ مدقق": "Choice A",
    "الخيار ب مدقق": "Choice B",
    "الخيار ج مدقق": "Choice C",
    "الخيار د مدقق": "Choice D",
    "الشرح مدقق": "Explanation",
    "معنى السؤال، وهدفه، وطريقة حله بشكل عام.": "Question Meaning",
}


def replace_placeholders_in_word(output_path, data):
    """Replaces placeholders in the Word document."""
    doc = DocxTemplate(WORD_MAIN_PATH)
    doc.render(data)
    doc.pics_to_replace
    doc.save(output_path)

    # for paragraph in doc.paragraphs:
    #     for key, value in PLACEHOLDERS.items():
    #         placeholder = f"[{value}]"
    #         if placeholder in paragraph.text:
    #             paragraph.text = paragraph.text.replace(placeholder, str(data.get(key, "")))


def replace_placeholders_in_powerpoint(prs, data):
    """Replaces placeholders in the PowerPoint presentation."""
    # for slide in prs.slides:
    #     for shape in slide.shapes:
    #         if shape.has_text_frame:
    #             text_frame = shape.text_frame
    #             for paragraph in text_frame.paragraphs:
    #                 for run in paragraph.runs:
    #                     for key, value in PLACEHOLDERS.items():
    #                         if f"[{value}]" in run.text:
    #                             run.text = run.text.replace(f"[{value}]", str(data.get(key, "")))
    for slide in prs.slides:
        for shape in slide.shapes:
            if shape.has_text_frame:
                for paragraph in shape.text_frame.paragraphs:
                    for run in paragraph.runs:
                        # Replace placeholders with actual content
                        for key, value in data.items():
                            if key in run.text:
                                run.text = run.text.replace(key, str(value))


def split_and_populate_explanation(doc, explanation_text, max_length=500):
    """Splits long explanations and populates a second Word template."""
    if len(explanation_text) > max_length:
        new_doc = docx.Document(WORD_MAIN_PATH)
        new_doc.paragraphs[0].text = explanation_text[max_length:]
        new_doc.save("Explanation_Part2.docx")
        doc.paragraphs[PLACEHOLDERS["Explanation"]].text = explanation_text[:max_length]
        return doc

    else:
        doc.paragraphs[PLACEHOLDERS["Explanation"]].text = explanation_text
        return doc


def generate_number():
    """Generate a 9-digit number."""
    digits = random.sample(range(10), 9)  # Ensure unique digits
    return "".join(map(str, digits))


def find_expression(number):
    """Find an expression that evaluates to 100."""
    digits = list(number)
    operators = ["+", "-", "*"]
    for perm in itertools.permutations(digits):  # Try all digit permutations
        for ops in itertools.product(
            operators, repeat=len(digits) - 1
        ):  # Try all operator combinations
            expression = ""
            for i, digit in enumerate(perm):
                expression += digit
                if i < len(perm) - 1:
                    expression += ops[i]

            try:
                if eval(expression) == 100:
                    return expression + " = 100"
            except (ZeroDivisionError, SyntaxError):
                pass  # Handle potential errors

    return "No solution found"  # If no solution is found


def get_student_name(telegram_id):
    """retrieves the student name."""
    result = database.get_data(
        "SELECT name FROM users WHERE telegram_id = ?", (telegram_id,)
    )

    if result:
        user_name = result[0]
    else:
        user_name = "Unknown User"

    return user_name


def get_student_phone_number(telegram_id):
    """retrieves the student phone number."""
    result = database.get_data(
        "SELECT name FROM users WHERE telegram_id = ?", (telegram_id,)
    )

    if result:
        user_phone = result[0]
    else:
        user_phone = "Unknown Phone"

    return user_phone
