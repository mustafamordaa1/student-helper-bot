import os
import docx
from docx.shared import Inches, Pt

from config import IMAGE_FOLDER, Q_AND_A_FILE_PATH


def q_and_a_document(df, saveto, data):
    """Updates the Q&A document with questions and answers."""
    try:
        doc = docx.Document(Q_AND_A_FILE_PATH)
    except FileNotFoundError:
        doc = docx.Document()

    for index, row in df.iterrows():
        question_text = row["نص السؤال مدقق"]
        choices = [
            f"A. {row['الخيار أ مدقق']}",
            f"B. {row['الخيار ب مدقق']}",
            f"C. {row['الخيار ج مدقق']}",
            f"D. {row['الخيار د مدقق']}",
        ]
        correct_answer = f"الجواب الصحيح: {row['الجواب الصحيح']}"
        explanation = row["الشرح مدقق"]

        doc.add_heading(f"Question {index+1}", level=2)
        doc.add_paragraph(question_text)

        if row["اسم الصورة"]:
            image_path = os.path.join(IMAGE_FOLDER, row["اسم الصورة"])
            try:
                doc.add_picture(image_path, width=Inches(4))
            except FileNotFoundError:
                doc.add_paragraph(f"Image not found: {image_path}")

        for choice in choices:
            doc.add_paragraph(choice)

        doc.add_paragraph(correct_answer)
        p = doc.add_paragraph()
        runner = p.add_run("Explanation:")
        runner.bold = True
        runner.font.size = Pt(12)
        doc.add_paragraph(explanation)
        doc.add_paragraph("")

    doc.save(Q_AND_A_FILE_PATH)
    print(f"Q&A document updated successfully at: {Q_AND_A_FILE_PATH}")
